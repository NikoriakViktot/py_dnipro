import os
import comtypes.client
import pandas as pd
from docx import Document


# Функція для конвертації файлу з .doc у .docx за допомогою comtypes
def convert_doc_to_docx(doc_path, docx_path):
    wdFormatXMLDocument = 12
    word = comtypes.client.CreateObject('Word.Application')
    doc = word.Documents.Open(doc_path)
    doc.SaveAs(docx_path, FileFormat=wdFormatXMLDocument)
    doc.Close()
    word.Quit()


# Функція для зчитування таблиці з документа Word
def read_word_table(file_path):
    doc = Document(file_path)
    print(doc)
    if len(doc.tables) == 0:
        raise ValueError("Документ не містить таблиць.")

    table = doc.tables[0]
    data = []
    for row in table.rows:
        text = [cell.text.strip() for cell in row.cells]
        data.append(text)

    return data


# Конвертація файлу .doc у .docx
input_file_path = r"C:\Users\user\PycharmProjects\py_dnipro\Витрати_11\42198_11.doc"
output_file_path = input_file_path.replace('.doc', '.docx')
convert_doc_to_docx(input_file_path, output_file_path)

# Зчитування даних з файлу .docx
data = read_word_table(output_file_path)

# Перевірка зчитаних даних
print("Зчитані дані з файлу .docx:")
for row in data[:5]:  # Виведення перших 5 рядків для перевірки
    print(row)

# Перетворення даних у DataFrame
columns = data[0]
rows = data[1:]
df = pd.DataFrame(rows, columns=columns)

# Видалення перших трьох рядків та останніх двох (зайві заголовки та підсумки)
df = df.drop(index=[0, 1, 2, -1, -2])

# Перетворення колонок на числовий тип, якщо це можливо
for col in df.columns[1:]:
    df[col] = pd.to_numeric(df[col], errors='coerce')

# Збереження у CSV
output_csv = "daily_water_discharge.csv"
df.to_csv(output_csv, index=False)

print(f"Дані збережено у файл: {output_csv}")
